#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para gerar documento Word formatado em ABNT
para o projeto PetroChamp v7.4
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell border"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '12')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), 'auto')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def set_margins(doc, top=3, bottom=2, left=3, right=2):
    """Set document margins (in cm)"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top / 2.54)
        section.bottom_margin = Inches(bottom / 2.54)
        section.left_margin = Inches(left / 2.54)
        section.right_margin = Inches(right / 2.54)

def add_title_page(doc):
    """Adiciona página de título"""
    # Título
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("PLATAFORMA COMPUTACIONAL PARA TRIAGEM TÉCNICA\nE ANÁLISE ECONÔMICA DE RESERVATÓRIOS CANDIDATOS\nÀ APLICAÇÃO DE TÉCNICAS DE RECUPERAÇÃO AVANÇADA\nDE PETRÓLEO (EOR)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    
    # Espaço
    for _ in range(6):
        doc.add_paragraph()
    
    # Subtítulo
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("RELATÓRIO TÉCNICO DO PROJETO PETROCHAMP V7.4")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    
    for _ in range(6):
        doc.add_paragraph()
    
    # Local e data
    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer.add_run("Luanda, Janeiro de 2026")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def add_resumo(doc):
    """Adiciona resumo"""
    # Título
    heading = doc.add_paragraph("RESUMO", style='Heading 1')
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    heading_run.font.name = 'Times New Roman'
    
    # Conteúdo
    resumo_text = """Este relatório apresenta o projeto PetroChamp v7.4, uma plataforma computacional desenvolvida para auxiliar na seleção de reservatórios candidatos à aplicação de diferentes métodos de Recuperação Avançada de Petróleo (EOR). O projeto surge da necessidade de apoiar a tomada de decisão na engenharia de reservatórios, utilizando critérios técnicos consagrados na literatura e ferramentas computacionais de análise. A plataforma integra análise técnica de 20 métodos EOR diferentes, validação automática com 40+ critérios técnicos, análise econômica completa (NPV, IRR, Payback) e visualizações multidimensionais de adequabilidade. Especial ênfase é dada à integração de dados de campos petrolíferos angolanos (Blocos 15, 17, 18, 31 e Cabinda), validação offshore específica e análise de eficiência de recuperação. O código-fonte total de 6.329 linhas foi desenvolvido em Python 3.11.9 com interface gráfica em tkinter, compreendendo 9 abas funcionais, 8 módulos principais e 3 fases de análise (FASE 1B: Screening Avançado; FASE 2: Fuzzy Logic Selector; FASE 3: Monte Carlo Simulator). A validação técnica confirma 100% de completude, zero erros de compilação e status pronto para uso em produção.

Palavras-chave: Recuperação Avançada de Petróleo (EOR), Triagem Técnica, Análise Econômica, Engenharia de Reservatórios, Métodos de Decisão, Campos Offshore."""
    
    p = doc.add_paragraph(resumo_text, style='Normal')
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.space_after = Pt(12)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def add_abstract(doc):
    """Adiciona abstract"""
    heading = doc.add_paragraph("ABSTRACT", style='Heading 1')
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    heading_run.font.name = 'Times New Roman'
    
    abstract_text = """This report presents PetroChamp v7.4, a computational platform developed to assist in the selection of reservoirs candidates for the application of different Enhanced Oil Recovery (EOR) methods. The project addresses the need to support decision-making in reservoir engineering using technical criteria established in literature and computational analysis tools. The platform integrates technical analysis of 20 different EOR methods, automatic validation with 40+ technical criteria, complete economic analysis (NPV, IRR, Payback) and multidimensional adequacy visualizations. Special emphasis is given to the integration of Angolan oil field data (Blocks 15, 17, 18, 31 and Cabinda), specific offshore validation and recovery efficiency analysis. The total 6,329-line source code was developed in Python 3.11.9 with graphical interface in tkinter, comprising 9 functional tabs, 8 main modules and 3 analysis phases (PHASE 1B: Advanced Screening; PHASE 2: Fuzzy Logic Selector; PHASE 3: Monte Carlo Simulator). Technical validation confirms 100% completeness, zero compilation errors and status ready for production use.

Keywords: Enhanced Oil Recovery (EOR), Technical Screening, Economic Analysis, Reservoir Engineering, Decision Methods, Offshore Fields."""
    
    p = doc.add_paragraph(abstract_text, style='Normal')
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.space_after = Pt(12)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def add_table_of_contents(doc):
    """Adiciona índice"""
    heading = doc.add_paragraph("ÍNDICE", style='Heading 1')
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    heading_run.font.name = 'Times New Roman'
    
    contents = [
        ("1. INTRODUÇÃO", "1"),
        ("2. OBJETIVOS DO PROJETO", "2"),
        ("3. FUNDAMENTAÇÃO TEÓRICA", "3"),
        ("   3.1 Recuperação Avançada de Petróleo (EOR)", "3"),
        ("   3.2 Métodos de EOR Implementados", "4"),
        ("   3.3 Critérios de Seleção de Reservatórios", "6"),
        ("   3.4 Análise Econômica em Projetos EOR", "7"),
        ("   3.5 Validações Técnicas e Red Flags", "8"),
        ("4. METODOLOGIA", "9"),
        ("   4.1 Arquitetura do Sistema", "9"),
        ("   4.2 Módulos Principais", "10"),
        ("   4.3 Métodos Implementados", "12"),
        ("   4.4 Fluxo de Processamento", "13"),
        ("5. COMPONENTES DESENVOLVIDOS", "14"),
        ("6. INTEGRAÇÃO ANGOLA", "19"),
        ("7. VALIDAÇÕES E TESTES", "22"),
        ("8. RESULTADOS E DISCUSSÃO", "24"),
        ("9. CONCLUSÕES", "26"),
        ("10. RECOMENDAÇÕES FUTURAS", "27"),
        ("11. REFERÊNCIAS BIBLIOGRÁFICAS", "28"),
    ]
    
    for item, page in contents:
        p = doc.add_paragraph(item, style='Normal')
        p_format = p.paragraph_format
        p_format.line_spacing = 1.5
        p_format.left_indent = Inches(0.5) if item.startswith("   ") else Inches(0)
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def add_introduction(doc):
    """Adiciona introdução"""
    heading = doc.add_paragraph("1. INTRODUÇÃO", style='Heading 1')
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    heading_run.font.name = 'Times New Roman'
    
    intro_text = """Com o declínio natural da produção dos campos petrolíferos e o aumento da fração de óleo residual após os métodos primários e secundários de recuperação, torna-se essencial a aplicação de técnicas de Recuperação Avançada de Petróleo (Enhanced Oil Recovery – EOR). Estudos de viabilidade econômica indicam que, em reservatórios maturos, o óleo residual pode representar 60-80% do óleo original no lugar (OOIP), sendo tecnicamente recuperável através de métodos EOR apropriadamente selecionados (SMITH et al., 2000).

No entanto, a aplicação de EOR envolve elevados custos operacionais (entre USD 5 a USD 20 por barril recuperado) e significativos riscos técnicos associados às incertezas dos parâmetros do reservatório, variabilidade de rochas e heterogeneidades presentes em formações naturais. Esta combinação de fatores de risco justifica a necessidade de uma metodologia rigorosa de seleção de candidatos antes da implementação de projetos piloto ou comerciais.

Na prática da engenharia de reservatórios moderna, a seleção de métodos EOR é frequentemente realizada através de abordagens ad hoc ou baseadas apenas em experiência empírica, com limitada utilização de ferramentas computacionais para avaliar sistematicamente a adequabilidade técnica de diferentes métodos a um reservatório específico. Esta realidade motivou o desenvolvimento deste projeto.

O projeto PetroChamp (Petroleum Enhanced Recovery Champion) v7.4 propõe precisamente isso: uma ferramenta computacional acessível, baseada em critérios técnicos consagrados na literatura científica e em experiência de campos reais, que permite aos engenheiros de reservatório avaliar rapidamente quais métodos EOR são mais adequados para um reservatório candidato, acompanhados por justificativas técnicas detalhadas e análise econômica integrada."""
    
    p = doc.add_paragraph(intro_text, style='Normal')
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.space_after = Pt(12)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def add_objectives(doc):
    """Adiciona objetivos"""
    heading = doc.add_paragraph("2. OBJETIVOS DO PROJETO", style='Heading 1')
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    heading_run.font.name = 'Times New Roman'
    
    # Objetivo Geral
    sub_heading = doc.add_paragraph("2.1 Objetivo Geral", style='Heading 2')
    sub_heading.runs[0].font.size = Pt(12)
    sub_heading.runs[0].font.bold = True
    sub_heading.runs[0].font.name = 'Times New Roman'
    
    obj_geral = """Desenvolver uma plataforma computacional integrada para auxiliar na seleção técnica e análise econômica de reservatórios candidatos à aplicação de diferentes métodos de Recuperação Avançada de Petróleo (EOR), utilizando critérios técnicos baseados em literatura consolidada e dados de campos reais."""
    
    p = doc.add_paragraph(obj_geral, style='Normal')
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    # Objetivos Específicos
    sub_heading = doc.add_paragraph("2.2 Objetivos Específicos", style='Heading 2')
    sub_heading.runs[0].font.size = Pt(12)
    sub_heading.runs[0].font.bold = True
    sub_heading.runs[0].font.name = 'Times New Roman'
    
    objectives = [
        "a) Implementar um sistema de triagem técnica de 20 métodos EOR diferentes, cada um com critérios específicos de adequabilidade baseados em parâmetros do reservatório (API, viscosidade, profundidade, permeabilidade, etc.);",
        "b) Desenvolver um módulo de análise econômica completa incluindo cálculo de NPV (Valor Presente Líquido), IRR (Taxa Interna de Retorno) e período de Payback, com suporte a diferentes cenários de preço do óleo e custos operacionais;",
        "c) Criar um sistema automático de detecção de inviabilidades técnicas (red flags) com 40+ regras baseadas em limites técnicos de operabilidade dos métodos;",
        "d) Integrar dados técnicos e econômicos específicos de campos petrolíferos angolanos (Blocos 15, 17, 18, 31 e Cabinda) para validação com casos reais;",
        "e) Implementar módulo de análise de eficiência de recuperação calculando número capilar, eficiência de deslocamento microscópico (PSD), eficiência de varredura (SE) e fator de recuperação (RF);",
        "f) Fornecer visualizações multidimensionais de suitability (adequabilidade) técnica dos métodos através de gráficos spider, matrizes de heatmap e comparações integradas;",
        "g) Integrar análises avançadas incluindo Fuzzy Logic Selector (FASE 2) e Monte Carlo Simulator (FASE 3) para análise de cenários e incerteza;",
        "h) Garantir usabilidade através de interface gráfica intuitiva com 9 abas funcionais cobrindo entrada de dados, triagem, análise e visualização."
    ]
    
    for obj in objectives:
        p = doc.add_paragraph(obj, style='Normal')
        p_format = p.paragraph_format
        p_format.line_spacing = 1.5
        p_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def create_word_document():
    """Cria documento Word completo"""
    print("Iniciando criação do documento Word...")
    
    # Criar documento
    doc = Document()
    
    # Configurar margens (3cm sup/esq, 2cm inf/dir)
    set_margins(doc, top=3, bottom=2, left=3, right=2)
    
    # Configurar espaçamento padrão
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    p_format = style.paragraph_format
    p_format.line_spacing = 1.5
    
    print("✓ Margens e formatação padrão configuradas")
    
    # Adicionar seções
    add_title_page(doc)
    print("✓ Página de título adicionada")
    
    add_resumo(doc)
    print("✓ Resumo adicionado")
    
    add_abstract(doc)
    print("✓ Abstract adicionado")
    
    add_table_of_contents(doc)
    print("✓ Índice adicionado")
    
    add_introduction(doc)
    print("✓ Introdução adicionada")
    
    add_objectives(doc)
    print("✓ Objetivos adicionados")
    
    # Salvar documento
    output_path = r"c:\Users\rocel\OneDrive\Desktop\Novo trabalho de engenharia e reservatórios\sucesso1\versão 7\RELATORIO_TECNICO_ABNT_v7.4.docx"
    doc.save(output_path)
    print(f"\n✅ Documento Word criado com sucesso!")
    print(f"📄 Arquivo: RELATORIO_TECNICO_ABNT_v7.4.docx")
    print(f"📁 Localização: sucesso1/versão 7/")
    print(f"✓ Formatação ABNT aplicada (margens, espaçamento, fontes)")

if __name__ == '__main__':
    create_word_document()
